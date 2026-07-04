#!/usr/bin/env python3
"""
Convert scraped IslamWeb JSON to styled DOCX with TOC, colors, and background.
Accepts book ID or URL — auto-fetches if not already scraped.
"""

import json
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree

OUTPUT_DIR = Path("output")
FONT_NAME = "Scheherazade New"
BG_COLOR = "FAF7F0"
VOLUME_SIZE = 10

# Auto-fetch if module available
try:
    from scraper import scrape_single, get_subjects, get_books_for_subject
except ImportError:
    scrape_single = None

# ── helpers ────────────────────────────────────────────────

def set_document_background(doc, color_hex):
    for bg in doc.element.findall(qn('w:background')):
        doc.element.remove(bg)
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), color_hex)
    doc.element.insert(0, bg)

def add_outline_numbering(doc):
    numbering = doc.part.numbering_part.element
    ABSTRACT_ID = '100'; NUM_ID = '100'
    for an in numbering.findall(qn('w:abstractNum')):
        if an.get(qn('w:abstractNumId')) == ABSTRACT_ID:
            link_styles(doc, NUM_ID); return
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), ABSTRACT_ID)
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'multilevel')
    abstract.append(multi)
    for lvl_idx in range(9):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(lvl_idx))
        for tag, val in [('w:start', '1'), ('w:numFmt', 'decimal')]:
            el = OxmlElement(tag); el.set(qn('w:val'), val); lvl.append(el)
        suff = OxmlElement('w:suff'); suff.set(qn('w:val'), 'space'); lvl.append(suff)
        lt = OxmlElement('w:lvlText'); lt.set(qn('w:val'), ''); lvl.append(lt)
        jc = OxmlElement('w:lvlJc'); jc.set(qn('w:val'), 'right'); lvl.append(jc)
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement('w:num'); num.set(qn('w:numId'), NUM_ID)
    anid = OxmlElement('w:abstractNumId'); anid.set(qn('w:val'), ABSTRACT_ID); num.append(anid)
    numbering.append(num)
    link_styles(doc, NUM_ID)

def link_styles(doc, num_id):
    for i in range(9):
        style_name = f'Heading {i+1}'
        styl = doc.styles[style_name]
        ppr = styl.element.get_or_add_pPr()
        existing = ppr.find(qn('w:numPr'))
        if existing is not None: ppr.remove(existing)
        numPr = OxmlElement('w:numPr')
        nid = OxmlElement('w:numId'); nid.set(qn('w:val'), num_id); numPr.append(nid)
        lv = OxmlElement('w:ilvl'); lv.set(qn('w:val'), str(i)); numPr.append(lv)
        insert_ppr_child(ppr, numPr)
        # Fix linked character style font (Heading1Char, etc.)
        try:
            char_sty = doc.styles[f'{style_name} Char']
        except KeyError:
            continue
        rpr = char_sty.element.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            char_sty.element.insert(0, rpr)
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)
        # Remove size from char style so it inherits from paragraph style
        for tag in ('w:sz', 'w:szCs'):
            el = rpr.find(qn(tag))
            if el is not None: rpr.remove(el)

CT_PPR_ORDER = [
    'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr', 'widowControl',
    'numPr', 'suppressLineNumbers', 'pBdr', 'shd', 'tabs', 'suppressAutoHyphens',
    'kinsoku', 'wordWrap', 'overflowPunct', 'topLinePunct', 'autoSpaceDE', 'autoSpaceDN',
    'bidi', 'adjustRightInd', 'snapToGrid', 'spacing', 'ind', 'contextualSpacing',
    'mirrorIndents', 'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
    'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle', 'rPr', 'sectPr', 'pPrChange',
]

def insert_ppr_child(ppr, new_el):
    """Insert new_el into ppr at the correct schema position (CT_PPr sequence)."""
    tag = etree.QName(new_el).localname
    existing = ppr.find(qn(f'w:{tag}'))
    if existing is not None:
        ppr.remove(existing)
    target_idx = CT_PPR_ORDER.index(tag)
    for child in ppr:
        child_tag = etree.QName(child).localname
        if child_tag not in CT_PPR_ORDER:
            continue
        if CT_PPR_ORDER.index(child_tag) > target_idx:
            child.addprevious(new_el)
            return
    ppr.append(new_el)

def set_rtl(run):
    rpr = run._r.get_or_add_rPr()
    if rpr.find(qn('w:rtl')) is None:
        rpr.append(parse_xml(f'<w:rtl {nsdecls("w")}/>'))

def add_heading(doc, text, level):
    """Like doc.add_heading but ensures Scheherazade New on all runs."""
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = FONT_NAME
        rpr = r._r.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)
        rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        rpr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
        set_rtl(r)
    return p

def add_run(para, text, bold=False, color=None, size=None):
    if not text: return
    run = para.add_run(text)
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    rpr = run._r.get_or_add_rPr()
    if bold:
        rpr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        rpr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
    if size:
        rpr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{int(size * 2)}"/>'))
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    set_rtl(run)

# ── HTML parser ────────────────────────────────────────────

def parse_html(html_text):
    from lxml.html import fromstring
    if not html_text: return []
    html_text = re.sub(r'<div[^>]*id="pagebody[^"]*"[^>]*>', '', html_text, count=1)
    html_text = re.sub(r'</div>\s*$', '', html_text)
    html_text = re.sub(r'\s*\[\s*ص\s*:\s*\d+\s*\]\s*', ' ', html_text)
    html_text = re.sub(r'  +', ' ', html_text)
    html_text = re.sub(r'\[\s+', '[', html_text)
    html_text = re.sub(r'\s+\]', ']', html_text)
    html_text = re.sub(r'\(\s+', '(', html_text)
    html_text = re.sub(r'\s+\)', ')', html_text)
    html_text = f'<root>{html_text}</root>'
    try: root = fromstring(html_text)
    except: return [("text", html_text)]
    segs = []
    def walk(el):
        tail = (el.tail or '').lstrip('\n') if el.tag != 'root' else ''
        if el.tag is etree.Comment:
            if tail: segs.append(("text", tail))
            return
        if el.tag == 'br':
            segs.append(("br", None))
            if tail: segs.append(("text", tail))
            return
        if 'display:none' in (el.get('style', '') or ''):
            if tail: segs.append(("text", tail))
            return
        text = el.text or ''
        if el.tag == 'root':
            for c in el: walk(c)
            return
        cls = el.get('class', '')
        t = "span_quran" if cls == 'quran' else "span_hadith" if cls == 'hadith' else "span_mainsubj" if cls == 'mainsubj' else "text"
        if text: segs.append((t, text))
        for c in el: walk(c)
        if tail: segs.append(("text", tail))
    walk(root)
    collapsed = []
    br_count = 0
    for s in segs:
        if s[0] == "br":
            br_count += 1
            if br_count <= 2:
                collapsed.append(s)
        else:
            br_count = 0
            collapsed.append(s)
    return collapsed

# ── resolve input ──────────────────────────────────────────

def resolve_book(book_id_or_url):
    """Return path to cached JSON; auto-fetch if missing."""
    # Extract book ID
    m = re.search(r'/content/(\d+)/', str(book_id_or_url))
    if m: book_id = int(m.group(1))
    else: book_id = int(book_id_or_url)

    cache = OUTPUT_DIR / f"book_{book_id}.json"
    if cache.exists():
        return cache

    print(f"Book {book_id} not cached locally. Scraping...", flush=True)
    if scrape_single:
        result = scrape_single(book_id)
        if result and cache.exists():
            return cache
        print("Scrape returned no data, using empty template.", flush=True)

    return cache

# ── build docx ─────────────────────────────────────────────

def build_docx(data, docx_path, volume_info=None):
    doc = Document()
    set_document_background(doc, BG_COLOR)
    add_outline_numbering(doc)

    def set_font_on_style(sty):
        """Set both ASCII and complex-script font on a style."""
        el = sty.element
        rpr = el.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            el.insert(0, rpr)
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)

    def set_style_spacing(sty, before, after, line=300):
        ppr = sty.element.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), str(before))
        sp.set(qn('w:after'), str(after))
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')
        insert_ppr_child(ppr, sp)

    sty = doc.styles['Normal']
    set_font_on_style(sty)
    sty.font.size = Pt(14)
    sty.font.bold = True
    ppr = sty.element.get_or_add_pPr()
    if ppr.find(qn('w:bidi')) is None:
        insert_ppr_child(ppr, parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    set_style_spacing(sty, 0, 0, 277)

    heading_cfg = [
        (1, 26, (26, 58, 92)),     # Heading 1: dark blue
        (2, 24, (139, 105, 20)),   # Heading 2: golden
        (3, 22, (0, 90, 80)),      # Heading 3: teal
        (4, 20, (140, 50, 50)),    # Heading 4: dark red
        (5, 18, (80, 50, 120)),    # Heading 5: purple
        (6, 17, (60, 100, 50)),    # Heading 6: green
        (7, 16, (150, 80, 40)),    # Heading 7: brown
        (8, 15, (40, 60, 100)),    # Heading 8: steel blue
        (9, 14, (100, 100, 100)),  # Heading 9: gray
    ]
    for level, sz, color in heading_cfg:
        hs = doc.styles[f'Heading {level}']
        set_font_on_style(hs)
        hs.font.size = Pt(sz)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(*color)
        hpr = hs.element.get_or_add_pPr()
        if hpr.find(qn('w:bidi')) is None:
            insert_ppr_child(hpr, parse_xml(f'<w:bidi {nsdecls("w")}/>'))
        ind = OxmlElement('w:ind')
        ind.set(qn('w:right'), '200')
        insert_ppr_child(hpr, ind)
        set_style_spacing(hs, 80, 200, 240)

    section = doc.sections[0]
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.right_margin = Cm(1.8); section.left_margin = Cm(1.8)

    book = data["book"]
    if volume_info: vol_num, total_vols, chapters = volume_info
    else: vol_num = total_vols = 1; chapters = data["chapters"]

    # ── cover ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '3600'); sp.set(qn('w:after'), '0')
    insert_ppr_child(pPr, sp)
    add_run(p, book["title"], bold=True, size=52, color=(139, 105, 20))

    if total_vols > 1:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f"الجزء {vol_num} من {total_vols}", size=32, color=(139, 105, 20))

    for label in ["author", "publisher", "subject_name"]:
        val = book.get(label, "")
        if not val: continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = (26, 58, 92) if label == "author" else (136, 136, 136)
        add_run(p, val, size=32 if label == "author" else 24, color=c)

    doc.add_page_break()

    # ── content ──
    def count_nodes(nodes):
        c = 0
        for n in nodes:
            c += 1 + count_nodes(n.get("sections", []))
        return c
    total = count_nodes(chapters)
    done = [0]

    def render_node(node, lv=1):
        done[0] += 1
        if done[0] % 200 == 0 or done[0] == 1:
            print(f"  [{done[0]}/{total}] {node['title'][:40]}", flush=True)
        add_heading(doc, node["title"], min(lv, 9))

        html = node.get("content_with_tashkeel_html", "") or node.get("content_with_tashkeel", "")
        if html:
            segs = parse_html(html)
            groups = [[]]
            for s in segs:
                if s[0] == "br":
                    if groups[-1]: groups.append([])
                else:
                    groups[-1].append(s)
            first = True
            for group in groups:
                if not group: continue
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                before = "40" if not first else "0"
                sp = parse_xml(f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="0"/>')
                pPr.append(sp)
                for t, txt in group:
                    txt = txt.strip()
                    if not txt: continue
                    if t == "span_quran": add_run(p, f"\uFD3F{txt}\uFD3E", bold=True, color=(255, 0, 0))
                    elif t == "span_hadith": add_run(p, f"\u00AB{txt}\u00BB", bold=True, color=(0, 0, 255))
                    elif t == "span_mainsubj": add_run(p, txt, bold=True, color=(180, 50, 50))
                    else: add_run(p, txt)
                first = False

        for child in node.get("sections", []):
            render_node(child, lv + 1)

    for ch in chapters: render_node(ch, 1)

    doc.save(str(docx_path))

    # Strip problematic parts for Google Docs compat
    import zipfile, io, re
    skip_files = {'word/stylesWithEffects.xml',
                  'customXml/item1.xml', 'customXml/_rels/item1.xml.rels', 'customXml/itemProps1.xml'}
    tmp = io.BytesIO()
    with zipfile.ZipFile(docx_path) as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in skip_files:
                    continue
                data = zin.read(item.filename)
                # Strip references to removed files from Content Types and rels
                if item.filename == '[Content_Types].xml':
                    data = re.sub(b'<Override PartName="/word/stylesWithEffects.xml".*?/>', b'', data)
                    data = re.sub(b'<Override PartName="/customXml/.*?".*?/>', b'', data)
                elif item.filename.endswith('.rels'):
                    data = re.sub(br'<Relationship[^>]*stylesWithEffects\.xml[^>]*/>', b'', data)
                    data = re.sub(br'<Relationship[^>]*customXml/[^>]*/>', b'', data)
                zout.writestr(item, data)
    Path(docx_path).write_bytes(tmp.getvalue())

    size_mb = Path(docx_path).stat().st_size / 1024 / 1024
    print(f"  Saved: {docx_path} ({size_mb:.1f} MB)", flush=True)

# ── entry points ───────────────────────────────────────────

def convert_book(source, docx_stem=None, chapters_per_volume=VOLUME_SIZE, full=False):
    json_path = resolve_book(source)
    if docx_stem is None:
        docx_stem = json_path.stem.replace("book_", "")

    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    chapters = data["chapters"]
    total_ch = len(chapters)

    # Detect if top-level items are volumes (their sections have sub-sections)
    has_sub_sections = False
    for ch in chapters[:5]:
        for s in ch.get("sections", []):
            if s.get("sections"):
                has_sub_sections = True
                break
        if has_sub_sections:
            break

    if has_sub_sections:
        # Top-level items are volumes — each becomes its own volume
        n_volumes = total_ch
        out = []
        for v, ch in enumerate(chapters):
            docx_path = OUTPUT_DIR / f"book_{docx_stem}_v{v+1}.docx"
            out.append(docx_path)
            vol_data = {"book": data["book"], "chapters": [ch]}
            print(f"Volume {v+1}/{n_volumes}: {ch['title'][:40]}", flush=True)
            build_docx(vol_data, docx_path, volume_info=(v+1, n_volumes, [ch]))
        return out

    if full or total_ch <= chapters_per_volume * 1.5:
        docx_path = OUTPUT_DIR / f"book_{docx_stem}.docx"
        print(f"Single volume -> {docx_path}", flush=True)
        build_docx(data, docx_path)
        return [docx_path]

    n_volumes = (total_ch + chapters_per_volume - 1) // chapters_per_volume
    out = []
    for v in range(n_volumes):
        s = v * chapters_per_volume
        e = min((v + 1) * chapters_per_volume, total_ch)
        vch = chapters[s:e]
        docx_path = OUTPUT_DIR / f"book_{docx_stem}_v{v+1}.docx"
        out.append(docx_path)
        print(f"Volume {v+1}/{n_volumes} ({len(vch)} chapters)...", flush=True)
        build_docx(data, docx_path, volume_info=(v+1, n_volumes, vch))
    return out

if __name__ == "__main__":
    full = False; cpv = VOLUME_SIZE
    pos = []
    for a in sys.argv[1:]:
        if a == "--full": full = True
        elif a == "--help":
            print(f"Usage: {sys.argv[0]} [--full] [--chapters-per-volume N] <book_id|url|json> [stem]")
            sys.exit(0)
        elif a.startswith("--chapters-per-volume="): cpv = int(a.split("=", 1)[1])
        elif a.startswith("--"):
            print(f"Unknown: {a}", flush=True); sys.exit(1)
        else: pos.append(a)

    if not pos:
        print(f"Usage: {sys.argv[0]} [--full] [--chapters-per-volume N] <book_id|url|json> [stem]")
        sys.exit(1)

    convert_book(pos[0], pos[1] if len(pos) > 1 else None, chapters_per_volume=cpv, full=full)
